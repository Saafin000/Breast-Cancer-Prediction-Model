"""
File Processing Service for handling various document and image formats
Supports PDF, Word, TXT, JPG, PNG, JPEG files
"""

import os
import cv2
import numpy as np
from PIL import Image, ImageEnhance, ImageFilter
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
import base64
from io import BytesIO

# Document processing
import PyPDF2
import fitz  # PyMuPDF
from docx import Document

# Optional textract import
try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

# Logging
from loguru import logger

class FileProcessor:
    def __init__(self):
        self.supported_image_formats = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']
        self.supported_document_formats = ['.pdf', '.docx', '.doc', '.txt']
        logger.info("File processor initialized")

    async def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process any supported file type and extract relevant information
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Dictionary containing processed file information
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {"error": f"File not found: {file_path}"}
            
            file_extension = file_path.suffix.lower()
            file_size = file_path.stat().st_size
            
            result = {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_extension": file_extension,
                "file_size_bytes": file_size,
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "processing_timestamp": str(np.datetime64('now'))
            }
            
            # Process based on file type
            if file_extension in self.supported_image_formats:
                image_result = await self._process_image(file_path)
                result.update(image_result)
                result["file_type"] = "image"
                
            elif file_extension in self.supported_document_formats:
                document_result = await self._process_document(file_path)
                result.update(document_result)
                result["file_type"] = "document"
                
            else:
                result["error"] = f"Unsupported file format: {file_extension}"
                result["supported_formats"] = {
                    "images": self.supported_image_formats,
                    "documents": self.supported_document_formats
                }
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return {"error": f"File processing failed: {str(e)}"}

    async def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process medical images and extract features"""
        try:
            # Load image with PIL
            pil_image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if pil_image.mode != 'RGB':
                pil_image = pil_image.convert('RGB')
            
            # Basic image information
            width, height = pil_image.size
            
            # Load image with OpenCV for analysis
            cv_image = cv2.imread(str(file_path))
            if cv_image is None:
                return {"error": "Could not load image with OpenCV"}
            
            # Image analysis
            image_analysis = await self._analyze_medical_image(cv_image, pil_image)
            
            # Convert image to base64 for potential API transmission
            buffered = BytesIO()
            pil_image.save(buffered, format="JPEG", quality=85)
            image_base64 = base64.b64encode(buffered.getvalue()).decode()
            
            return {
                "image_analysis": image_analysis,
                "image_properties": {
                    "width": width,
                    "height": height,
                    "channels": len(pil_image.getbands()),
                    "mode": pil_image.mode,
                    "format": pil_image.format
                },
                "image_base64": image_base64[:100] + "..." if len(image_base64) > 100 else image_base64,  # Truncated for display
                "full_image_available": True
            }
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}")
            return {"error": f"Image processing failed: {str(e)}"}

    async def _analyze_medical_image(self, cv_image: np.ndarray, pil_image: Image.Image) -> Dict[str, Any]:
        """Analyze medical images for tumor characteristics"""
        try:
            # Convert to grayscale for analysis
            gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
            
            # Basic image statistics
            mean_intensity = np.mean(gray)
            std_intensity = np.std(gray)
            
            # Edge detection for border irregularity
            edges = cv2.Canny(gray, 50, 150)
            edge_density = np.sum(edges > 0) / (edges.shape[0] * edges.shape[1])
            
            # Contour analysis for shape characteristics
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            contour_analysis = {}
            if contours:
                # Find largest contour (assume it's the main tumor)
                largest_contour = max(contours, key=cv2.contourArea)
                
                # Calculate contour properties
                area = cv2.contourArea(largest_contour)
                perimeter = cv2.arcLength(largest_contour, True)
                
                # Compactness (circularity)
                if perimeter > 0:
                    compactness = (perimeter ** 2) / (4 * np.pi * area) if area > 0 else 0
                else:
                    compactness = 0
                
                # Convex hull for concavity analysis
                hull = cv2.convexHull(largest_contour)
                hull_area = cv2.contourArea(hull)
                solidity = area / hull_area if hull_area > 0 else 0
                
                contour_analysis = {
                    "tumor_area_pixels": float(area),
                    "tumor_perimeter_pixels": float(perimeter),
                    "compactness_estimate": float(compactness),
                    "solidity": float(solidity),
                    "irregularity_score": float(1 - solidity),  # Higher = more irregular
                    "contour_count": len(contours)
                }
            
            # Texture analysis using Local Binary Pattern concept
            texture_variance = cv2.Laplacian(gray, cv2.CV_64F).var()
            
            return {
                "basic_statistics": {
                    "mean_intensity": float(mean_intensity),
                    "std_intensity": float(std_intensity),
                    "intensity_range": [float(gray.min()), float(gray.max())]
                },
                "edge_analysis": {
                    "edge_density": float(edge_density),
                    "border_irregularity": "high" if edge_density > 0.1 else "low"
                },
                "texture_analysis": {
                    "texture_variance": float(texture_variance),
                    "texture_level": "high" if texture_variance > 500 else "moderate" if texture_variance > 100 else "low"
                },
                "contour_analysis": contour_analysis,
                "ai_interpretation": self._interpret_image_features(mean_intensity, edge_density, texture_variance, contour_analysis)
            }
            
        except Exception as e:
            logger.error(f"Error in medical image analysis: {str(e)}")
            return {"error": f"Image analysis failed: {str(e)}"}

    def _interpret_image_features(self, mean_intensity: float, edge_density: float, texture_variance: float, contour_analysis: Dict) -> Dict[str, str]:
        """Interpret image features for medical context"""
        interpretations = []
        
        # Intensity interpretation
        if mean_intensity < 50:
            interpretations.append("Low intensity regions may indicate dense tissue")
        elif mean_intensity > 200:
            interpretations.append("High intensity regions may indicate less dense areas")
        
        # Edge interpretation
        if edge_density > 0.15:
            interpretations.append("High edge density suggests irregular borders - concerning for malignancy")
        elif edge_density < 0.05:
            interpretations.append("Low edge density suggests smooth borders - favorable for benign lesions")
        
        # Texture interpretation
        if texture_variance > 1000:
            interpretations.append("High texture variance indicates heterogeneous tissue - may suggest malignancy")
        elif texture_variance < 100:
            interpretations.append("Low texture variance indicates homogeneous tissue - favorable for benign lesions")
        
        # Contour interpretation
        if contour_analysis:
            irregularity = contour_analysis.get("irregularity_score", 0)
            if irregularity > 0.3:
                interpretations.append("Irregular shape characteristics - concerning pattern")
            elif irregularity < 0.1:
                interpretations.append("Regular shape characteristics - favorable pattern")
        
        return {
            "summary": "; ".join(interpretations) if interpretations else "Standard image characteristics",
            "recommendation": "Professional radiological interpretation required for definitive diagnosis"
        }

    async def _process_document(self, file_path: Path) -> Dict[str, Any]:
        """Process text documents and extract medical information"""
        try:
            file_extension = file_path.suffix.lower()
            extracted_text = ""
            
            # Extract text based on file type
            if file_extension == '.pdf':
                extracted_text = self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                extracted_text = self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()
            else:
                # Try textract for other formats if available
                if TEXTRACT_AVAILABLE:
                    try:
                        extracted_text = textract.process(str(file_path)).decode('utf-8')
                    except:
                        return {"error": f"Could not extract text from {file_extension} file"}
                else:
                    return {"error": f"Unsupported file format: {file_extension}. Install textract for broader support."}
            
            # Analyze extracted text
            text_analysis = await self._analyze_medical_text(extracted_text)
            
            return {
                "extracted_text": extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text,
                "full_text_length": len(extracted_text),
                "text_analysis": text_analysis,
                "medical_entities": self._extract_medical_entities(extracted_text)
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return {"error": f"Document processing failed: {str(e)}"}

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Try PyMuPDF first (better for complex PDFs)
        try:
            doc = fitz.open(str(file_path))
            for page in doc:
                text += page.get_text()
            doc.close()
            if text.strip():
                return text
        except:
            pass
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except:
            pass
        
        return text

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from Word documents"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""

    async def _analyze_medical_text(self, text: str) -> Dict[str, Any]:
        """Analyze extracted text for medical relevance"""
        try:
            # Basic text statistics
            word_count = len(text.split())
            char_count = len(text)
            line_count = len(text.split('\n'))
            
            # Look for medical keywords
            medical_keywords = [
                'tumor', 'cancer', 'malignant', 'benign', 'breast', 'mammogram',
                'biopsy', 'ultrasound', 'mri', 'ct scan', 'radiology', 'oncology',
                'diagnosis', 'pathology', 'metastasis', 'lesion', 'mass', 'lump',
                'fibroid', 'cyst', 'calcification', 'density', 'texture', 'border'
            ]
            
            found_keywords = []
            text_lower = text.lower()
            
            for keyword in medical_keywords:
                if keyword in text_lower:
                    count = text_lower.count(keyword)
                    found_keywords.append({"keyword": keyword, "count": count})
            
            # Extract potential measurements
            measurements = self._extract_measurements(text)
            
            # Determine document relevance
            relevance_score = len(found_keywords) / len(medical_keywords)
            
            if relevance_score > 0.3:
                relevance = "high"
            elif relevance_score > 0.1:
                relevance = "moderate"
            else:
                relevance = "low"
            
            return {
                "text_statistics": {
                    "word_count": word_count,
                    "character_count": char_count,
                    "line_count": line_count
                },
                "medical_relevance": {
                    "relevance_level": relevance,
                    "relevance_score": relevance_score,
                    "keywords_found": found_keywords[:10]  # Top 10
                },
                "extracted_measurements": measurements,
                "summary": self._generate_text_summary(text, found_keywords)
            }
            
        except Exception as e:
            logger.error(f"Error analyzing medical text: {str(e)}")
            return {"error": f"Text analysis failed: {str(e)}"}

    def _extract_measurements(self, text: str) -> List[Dict[str, Any]]:
        """Extract numerical measurements from text"""
        import re
        
        measurements = []
        
        # Common measurement patterns
        patterns = [
            (r'(\d+\.?\d*)\s*(cm|mm|inches?)', 'size'),
            (r'(\d+\.?\d*)\s*(ml|cc)', 'volume'), 
            (r'area[:\s]*(\d+\.?\d*)', 'area'),
            (r'radius[:\s]*(\d+\.?\d*)', 'radius'),
            (r'diameter[:\s]*(\d+\.?\d*)', 'diameter'),
            (r'(\d+\.?\d*)\s*x\s*(\d+\.?\d*)', 'dimensions')
        ]
        
        for pattern, measurement_type in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    if measurement_type == 'dimensions':
                        measurements.append({
                            "type": measurement_type,
                            "value": f"{match[0]} x {match[1]}",
                            "unit": "units"
                        })
                    else:
                        measurements.append({
                            "type": measurement_type,
                            "value": float(match[0]),
                            "unit": match[1] if len(match) > 1 else "units"
                        })
                else:
                    measurements.append({
                        "type": measurement_type,
                        "value": float(match),
                        "unit": "units"
                    })
        
        return measurements[:5]  # Limit to 5 measurements

    def _generate_text_summary(self, text: str, keywords: List[Dict]) -> str:
        """Generate a summary of the document content"""
        if len(text) < 100:
            return text
        
        # Simple extractive summarization
        sentences = text.split('.')
        relevant_sentences = []
        
        keyword_list = [kw["keyword"] for kw in keywords]
        
        for sentence in sentences[:20]:  # Check first 20 sentences
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in keyword_list):
                relevant_sentences.append(sentence.strip())
                if len(relevant_sentences) >= 3:
                    break
        
        if relevant_sentences:
            return ". ".join(relevant_sentences) + "."
        else:
            # Return first few sentences if no keywords found
            return ". ".join(sentences[:2]) + "."

    def _extract_medical_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract medical entities from text"""
        import re
        
        entities = {
            "diagnoses": [],
            "procedures": [],
            "measurements": [],
            "medications": [],
            "anatomy": []
        }
        
        # Define entity patterns
        entity_patterns = {
            "diagnoses": [
                r'diagnosed?\s+with\s+([^.]+)',
                r'(benign|malignant)\s+([^.]+)',
                r'(carcinoma|adenocarcinoma|sarcoma|lymphoma)',
                r'(stage\s+[IVX1-4]+)'
            ],
            "procedures": [
                r'(biopsy|mammogram|ultrasound|mri|ct\s+scan)',
                r'(surgery|mastectomy|lumpectomy)',
                r'(chemotherapy|radiation|immunotherapy)'
            ],
            "anatomy": [
                r'(breast|chest|thorax|axilla|lymph\s+node)',
                r'(left|right)\s+(breast|side)',
                r'(upper|lower)\s+(quadrant|pole)'
            ]
        }
        
        text_lower = text.lower()
        
        for entity_type, patterns in entity_patterns.items():
            for pattern in patterns:
                matches = re.findall(pattern, text_lower, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, tuple):
                        entity_text = " ".join(match)
                    else:
                        entity_text = match
                    
                    if entity_text not in entities[entity_type]:
                        entities[entity_type].append(entity_text)
        
        # Clean up empty lists
        entities = {k: v for k, v in entities.items() if v}
        
        return entities

    async def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported file formats"""
        return {
            "images": self.supported_image_formats,
            "documents": self.supported_document_formats,
            "description": {
                "images": "Medical images for visual analysis",
                "documents": "Medical reports, research papers, clinical notes"
            }
        }
